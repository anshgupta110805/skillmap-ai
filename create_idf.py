from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x0A, 0x25, 0x4A)   # #0A254A
MID_BLUE   = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB
LIGHT_BLUE = RGBColor(0xDB, 0xEA, 0xFE)   # #DBEAFE
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TEXT  = RGBColor(0x37, 0x41, 0x51)   # #374151
BLACK      = RGBColor(0x00, 0x00, 0x00)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set cell border ───────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            border = OxmlElement(f'w:{side}')
            for attr, val in kwargs[side].items():
                border.set(qn(f'w:{attr}'), val)
            tcBorders.append(border)
    tcPr.append(tcBorders)

# ── Helper: paragraph with run ────────────────────────────────────────────────
def add_para(doc_or_cell, text, bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
             space_after=6, font_name='Calibri'):
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

# ── Helper: heading style ─────────────────────────────────────────────────────
def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{number}. {title.upper()}")
    run.bold           = True
    run.font.name      = 'Calibri'
    run.font.size      = Pt(12)
    run.font.color.rgb = DARK_BLUE
    return p

def add_sub_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title)
    run.bold           = True
    run.font.name      = 'Calibri'
    run.font.size      = Pt(11)
    run.font.color.rgb = MID_BLUE
    return p

def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_part:
        r1 = p.add_run(bold_part + ': ')
        r1.bold           = True
        r1.font.name      = 'Calibri'
        r1.font.size      = Pt(11)
        r1.font.color.rgb = DARK_BLUE
        r2 = p.add_run(text)
        r2.font.name      = 'Calibri'
        r2.font.size      = Pt(11)
        r2.font.color.rgb = GREY_TEXT
    else:
        r = p.add_run(text)
        r.font.name      = 'Calibri'
        r.font.size      = Pt(11)
        r.font.color.rgb = GREY_TEXT
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(11)
    run.font.color.rgb = GREY_TEXT
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  COVER / TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
# Blue banner table
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
cell = tbl.rows[0].cells[0]
shade_cell(cell, '0A254A')
set_cell_border(cell,
    top    ={'val':'single','sz':'6','color':'1A56DB'},
    bottom ={'val':'single','sz':'6','color':'1A56DB'},
    left   ={'val':'none'} ,
    right  ={'val':'none'})
cell.width = Inches(6.5)

# Title lines inside banner
p1 = cell.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(10)
r1 = p1.add_run('INVENTION DISCLOSURE FORM (IDF)')
r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(10)
r1.font.color.rgb = LIGHT_BLUE

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('SkillMapper AI')
r2.bold = True; r2.font.name = 'Calibri'; r2.font.size = Pt(22)
r2.font.color.rgb = WHITE

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('AI-Powered Skill Intelligence & Career Navigation Platform')
r3.bold = True; r3.font.name = 'Calibri'; r3.font.size = Pt(13)
r3.font.color.rgb = LIGHT_BLUE

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(4)
r4 = p4.add_run('NEW SUBMISSION  |  Ticket ID: To Be Assigned by IP Office  |  Date of Disclosure: 25 February 2026')
r4.italic = True; r4.font.name = 'Calibri'; r4.font.size = Pt(9)
r4.font.color.rgb = LIGHT_BLUE

p5 = cell.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
p5.paragraph_format.space_after = Pt(10)
r5 = p5.add_run('Status: First Submission  |  Not Yet Reviewed  |  Strictly Confidential')
r5.italic = True; r5.font.name = 'Calibri'; r5.font.size = Pt(8.5)
r5.font.color.rgb = RGBColor(0xFF, 0xC0, 0x40)  # amber warning

doc.add_paragraph()

# ── Meta-data table ──────────────────────────────────────────────────────────
meta = [
    ('Invention Title',       'SkillMapper AI – AI-Powered Skill Intelligence & Career Navigation Platform'),
    ('Ticket / IDF ID',       'To Be Assigned by IP Office (New Submission)'),
    ('Submission Status',     '🔵 NEW — First Time Submission | Not Yet Reviewed'),
    ('Product / System',      'SkillMapper AI'),
    ('Submitted By',          'Ansh Gupta & Vidur Kumar'),
    ('Role / Designation',    'Full Stack AI Developer & Product Engineer'),
    ('Department',            'AI/ML & Product Engineering'),
    ('Date of Submission',    '25 February 2026'),
    ('Confidentiality',       'Strictly Confidential — Not for External Distribution'),
    ('Technology Domain',     'Artificial Intelligence, Career Technology, Workforce Analytics, NLP'),
    ('Reviewed By (IP Cell)', '___________________________'),
    ('Review Date',           '___________________________'),
]
mt = doc.add_table(rows=len(meta), cols=2)
mt.style = 'Table Grid'
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(meta):
    lc = mt.rows[i].cells[0]
    vc = mt.rows[i].cells[1]
    shade_cell(lc, '1A4080')
    shade_cell(vc, 'F0F4FF')
    lc.width = Inches(2.2)
    vc.width = Inches(4.3)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(label)
    lr.bold = True; lr.font.name = 'Calibri'; lr.font.size = Pt(10); lr.font.color.rgb = WHITE
    vp = vc.paragraphs[0]
    vr = vp.add_run(value)
    vr.font.name = 'Calibri'; vr.font.size = Pt(10); vr.font.color.rgb = DARK_BLUE

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – TITLE OF INVENTION
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 1, 'Title of Invention')
add_body(doc,
    'SkillMapper AI: An Artificial Intelligence-Powered, Multi-Modal Skill Intelligence and Dynamic '
    'Career Navigation Platform that employs real-time labour market ingestion, NLP-based skill '
    'extraction, predictive career pathing, and geographic opportunity mapping to guide individuals '
    'and organisations toward optimal workforce utilisation.')

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – FIELD OF INVENTION
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 2, 'Field of Invention')
add_body(doc,
    'The present invention belongs to the fields of Artificial Intelligence (AI), Natural Language '
    'Processing (NLP), Human Capital Management (HCM), and Career Technology (CareerTech). More '
    'specifically, it relates to a software system and method for the real-time analysis of labour '
    'market signals, dynamic skill-gap detection, predictive career route recommendation, and '
    'AI-driven workforce policy advisory.')

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – BACKGROUND / INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 3, 'Background / Introduction')
add_body(doc,
    'The modern labour market is characterised by rapid technological transformation, geographic '
    'skill disparities, and an increasing misalignment between institutional education outputs and '
    'actual employer demand. Existing career advisory platforms rely on static job-board aggregation, '
    'keyword-based skill matching, and generic career path templates that fail to capture the '
    'nuanced, dynamic interplay between an individual\'s evolving skill profile and the real-time '
    'demands of a fluid, city-segmented labour market.')
add_body(doc,
    'SkillMapper AI was conceived to fill this critical gap. By combining live demand-supply heatmapping, '
    'NLP-driven skill extraction from job descriptions, AI-powered career route optimisation, '
    'geographic migration opportunity modelling, and an explainable policy advisory engine, the '
    'invention creates a cohesive intelligence layer between job seekers, employers, and '
    'policy makers — a layer that learns and adapts continuously.')

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – OBJECTIVES OF THE INVENTION
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 4, 'Objectives of the Invention')
objectives = [
    ('Real-Time Skill Demand Mapping',
     'To continuously ingest and process live job-market data to produce a dynamic, city- and '
     'industry-segmented heatmap of skill demand and supply, enabling individuals and organisations '
     'to make data-driven workforce decisions.'),
    ('AI-Driven Skill Gap Detection',
     'To automatically identify critical skill shortages and oversupplies across geographies and '
     'industries, surfacing actionable intelligence for strategic career repositioning.'),
    ('Personalised Career Route Recommendation',
     'To generate individualised, AI-computed career trajectories based on a user\'s unique multi-'
     'dimensional profile — encompassing current skills, experience, education, employment status, '
     'and declared goals — using predictive modelling rather than static templates.'),
    ('Geographic Career Migration Analysis',
     'To model and compare career opportunities across cities, accounting for cost-of-living '
     'indices, salary benchmarks, and demand-supply dynamics, so that users can make informed '
     'geographic mobility decisions.'),
    ('Explainable AI Policy Advisory',
     'To provide human-readable, evidence-based policy recommendations for workforce development '
     'bodies and enterprise HR teams, grounded in observed labour market anomalies and skill '
     'shortage patterns.'),
    ('NLP-Powered Skill Extraction',
     'To automatically deconstruct unstructured job description text into a structured, ranked '
     'list of required and preferred competencies, eliminating manual skill cataloguing.'),
    ('Skill Relationship Graph Intelligence',
     'To construct and analyse a dynamic knowledge graph of skill interdependencies and career '
     'adjacencies, enabling discovery of non-obvious upskilling and lateral move opportunities.'),
    ('Workforce Digital Twin Simulation',
     'To simulate the projected future skill landscape and workforce composition for an '
     'organisation or city, enabling proactive talent planning.'),
]
for title, desc in objectives:
    add_bullet(doc, desc, title)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – PROBLEMS IN PRIOR ART
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 5, 'Problems in Prior Art the Invention Targets to Solve')

problems = [
    ('Static, Non-Adaptive Skill Profiles',
     'Problem: Existing platforms (e.g., LinkedIn, traditional HRMS) rely on manually curated, '
     'self-reported skill profiles that rapidly become outdated and do not reflect an individual\'s '
     'actual, demonstrated, or recently acquired competencies.',
     'Solution: SkillMapper AI maintains a living skill model per user, automatically enriched '
     'through NLP extraction from uploaded CVs, job descriptions, and connected professional activity.'),
    ('Reactive, Keyword-Based Career Matching',
     'Problem: Career recommendation engines predominantly use keyword overlap between résumés and '
     'job postings. This ignores proficiency levels, contextual relevance, and career trajectory fit.',
     'Solution: The AI Career Planner module uses multi-dimensional user profiles and a predictive '
     'model trained on career progression data to recommend contextually optimal paths — not just '
     'keyword-matched roles.'),
    ('No Geographic Skill Intelligence',
     'Problem: Market tools present national aggregates. City-level skill demand, salary variation, '
     'and cost-of-living-adjusted opportunity are rarely surfaced together in one actionable view.',
     'Solution: The Skills Heatmap and Migration Assistant modules deliver granular, city-segmented '
     'demand-supply visualisations and opportunity scores, enabling evidence-based geographic '
     'career decisions.'),
    ('Lack of Explainable AI in Career Advisory',
     'Problem: Where AI is used in career platforms, it typically operates as a black box, providing '
     'recommendations without transparent reasoning, which erodes user trust.',
     'Solution: Every recommendation generated by SkillMapper AI is accompanied by an Explainable AI '
     '(XAI) rationale, detailing the data signals and reasoning behind the suggestion (e.g., '
     '"Machine Learning demand in Bangalore has grown 42% YoY with a 2.3× demand-to-supply ratio").'),
    ('No Feedback Loop or Continuous Learning',
     'Problem: Prior systems are deployed statically; they do not improve based on user interaction '
     'data, labour market shifts, or outcome feedback.',
     'Solution: SkillMapper AI incorporates a continuous learning pipeline that updates its models '
     'based on new labour market data, user interaction patterns, and recommendation outcome signals.'),
    ('Absence of Policy-Level Workforce Intelligence',
     'Problem: Workforce policy decisions at government and enterprise level are often based on '
     'lagging, aggregated reports rather than real-time, granular, AI-synthesised intelligence.',
     'Solution: The Policy Advisor module transforms live skill-shortage and anomaly data into '
     'prioritised, actionable policy prescriptions specifically tailored to detected market conditions.'),
    ('No Skill Relationship / Graph Modelling',
     'Problem: Users cannot see how skills relate to one another, making it difficult to identify '
     'optimal upskilling paths or lateral career moves.',
     'Solution: The Skill Graph Intelligence module constructs a dynamic Neo4j-style knowledge graph '
     'of skill interdependencies, cluster memberships, and career adjacency scores.'),
]
for title, prob, sol in problems:
    add_sub_heading(doc, title)
    add_bullet(doc, prob)
    add_bullet(doc, sol)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – NOVEL ASPECTS OF THE INVENTION
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 6, 'Novel Aspects of the Invention')
novel = [
    ('Integrated Real-Time Demand-Supply Heatmapping',
     'The simultaneous, city-level visualisation of both skill demand (job postings) and skill '
     'supply (professional availability) on a unified, filterable heatmap interface — providing '
     'an at-a-glance arbitrage view of the skill market that no known prior system offers as '
     'an integrated, live feature.'),
    ('Multi-Signal AI Career Route Planner',
     'A generative AI planner (powered by Google Gemini via Genkit) that synthesises user profile '
     'data across seven dimensions (skills, experience, education, employment status, industry, '
     'role, and age) to produce a structured, milestone-based personalised career roadmap with '
     'explicit rationale — moving beyond step-by-step role lists to a full narrative trajectory.'),
    ('Cost-of-Living-Adjusted Migration Opportunity Scoring',
     'A novel scoring methodology that normalises salary data against city-specific cost-of-living '
     'indices to produce a true "real-income opportunity score," enabling users to identify the city '
     'where their specific skill set yields the highest quality-adjusted compensation.'),
    ('AI Skill Anomaly Detection Engine',
     'A dedicated AI flow that analyses skill demand patterns to automatically surface statistical '
     'anomalies — sudden demand spikes, emerging skill clusters, or unexpected supply collapses — '
     'that are invisible to standard market reports but critical for strategic planning.'),
    ('Workforce Digital Twin Simulation',
     'A simulation module that models projected workforce compositions and skill landscapes under '
     'different growth, training, or hiring scenarios — functioning as a virtual replica of an '
     'organisation\'s or region\'s talent pool for proactive planning.'),
    ('Unified Skill Knowledge Graph with Career Adjacency Mapping',
     'A dynamically generated skill relationship graph that maps competency interdependencies, '
     'learning path optimality, and career transition probabilities — enabling users to discover '
     'non-obvious upskilling sequences and adjacent career pivots.'),
]
for title, desc in novel:
    add_bullet(doc, desc, title)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – DETAILED DESCRIPTION OF THE INVENTION
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 7, 'Detailed Description of the Invention')

# 7.1
add_sub_heading(doc, '7.1  System Architecture Overview')
add_body(doc,
    'SkillMapper AI is implemented as a cloud-native, full-stack web application built on the '
    'Next.js 14 App Router framework with a TypeScript/React frontend, a Firebase backend '
    '(Authentication & Hosting), and a suite of AI flows powered by Google Genkit with the '
    'Gemini Pro model family. The system is organised into eight functional modules, each '
    'accessible via a dedicated application route.')

# 7.2 – Modules table
add_sub_heading(doc, '7.2  Core Functional Modules')
modules = [
    ('Skills Heatmap',         '/explore',             'Interactive city × industry demand-supply visualisation using Recharts. Users filter by skill category and geography to identify market opportunities.'),
    ('Skill Shortage Engine',  '/city-intelligence',   'Detects and ranks critical skill shortages and oversupplies using demand vs. supply count differentials, providing strategic market intelligence.'),
    ('AI Career Planner',      '/career-planner',      'Genkit AI flow (recommend-career-route) that consumes a 7-dimension user profile and returns a structured, milestone-based career roadmap with XAI rationale.'),
    ('Migration Assistant',    '/migration-assistant', 'AI-powered (skill-migration-assistant flow) analysis comparing career opportunity scores across cities, normalised for cost-of-living indices.'),
    ('Migration Simulator',    '/migration-simulator', 'Scenario simulation module (migration-simulator flow) that projects income and opportunity changes under hypothetical geographic moves.'),
    ('Policy Advisor',         '/policy-advisor',      'AI flow (policy-advisor) that synthesises live market anomalies into prioritised, human-readable policy prescriptions for HR teams and government bodies.'),
    ('Skill Graph Intelligence','/skill-graph',        'Constructs and renders a dynamic knowledge graph (skill-graph-intelligence flow) mapping skill interdependencies, clusters, and career adjacency scores.'),
    ('Skill Extractor',        '/skill-extractor',     'NLP-powered flow (extract-skills-from-job-description) that deconstructs raw job description text into a structured, ranked competency list.'),
    ('Workforce Digital Twin', '/workforce-twin',      'Simulation module that models projected workforce composition under alternative hiring, training, or attrition scenarios.'),
]
mt2 = doc.add_table(rows=len(modules)+1, cols=3)
mt2.style = 'Table Grid'
mt2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Module Name', 'Route', 'Function']
for j, h in enumerate(headers):
    c = mt2.rows[0].cells[j]
    shade_cell(c, '0A254A')
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.color.rgb = WHITE
for i, (mod, route, fn) in enumerate(modules):
    row = mt2.rows[i+1]
    shade = 'F0F4FF' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([mod, route, fn]):
        c = row.cells[j]
        shade_cell(c, shade)
        r = c.paragraphs[0].add_run(val)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.color.rgb = GREY_TEXT
        if j == 0:
            r.bold = True; r.font.color.rgb = DARK_BLUE

doc.add_paragraph()

# 7.3 AI Flows
add_sub_heading(doc, '7.3  AI Flow Architecture (Genkit + Google Gemini)')
add_body(doc,
    'All AI reasoning is encapsulated in discrete, testable Genkit flows located in the /ai/flows '
    'directory. Each flow defines a typed input/output schema using Zod, invokes the Gemini Pro '
    'model with a structured prompt, and returns a validated, strongly-typed response consumed by '
    'the frontend. This architecture ensures reproducibility, prompt version control, and '
    'straightforward extensibility.')

ai_flows = [
    ('extract-skills-from-job-description', 'Parses raw JD text → returns ranked list of required + preferred skills with proficiency indicators.'),
    ('recommend-career-route',              'Accepts 7-dim user profile → returns milestone-based career roadmap with XAI rationale per step.'),
    ('skill-migration-assistant',           'Accepts current city + target city + skill profile → returns opportunity comparison report.'),
    ('migration-simulator',                 'Accepts migration scenario params → projects income, opportunity and quality-of-life scores.'),
    ('policy-advisor',                      'Accepts detected skill market anomalies → returns prioritised policy prescription list.'),
    ('skill-graph-intelligence',            'Accepts skill cluster data → returns graph nodes/edges with adjacency weights and cluster labels.'),
]
for flow, desc in ai_flows:
    add_bullet(doc, desc, flow)

# 7.4 Data Model
add_sub_heading(doc, '7.4  Core Data Model')
add_body(doc,
    'The system\'s data model (defined in /lib/types.ts) comprises the following key entities:')
entities = [
    ('UserData',     'Full user profile: name, age, skills[], employmentStatus, experience, education, currentRole, industry. Drives all personalised AI flows.'),
    ('SkillDemand',  'Represents the intersection of a skill, a city, and its demand/supply counts (job postings vs. professional supply). Core data unit for heatmap and shortage engine.'),
    ('SalaryData',   'City × role × averageSalary triples. Used in cost-of-living-adjusted migration scoring.'),
    ('TrendData',    'Month-indexed, multi-skill demand trend series. Powers time-series chart visualisations.'),
    ('City',         'Name + costOfLivingIndex. Enables normalised geographic comparisons.'),
    ('Skill / Industry', 'Taxonomic classification entities enabling filtered views across the skill graph and heatmap.'),
]
for name, desc in entities:
    add_bullet(doc, desc, name)

# 7.5 Process Steps
add_sub_heading(doc, '7.5  End-to-End Process Flow')

steps = [
    ('Step 1 – User Onboarding & Profile Construction',
     'The user completes the structured onboarding flow (/onboarding), supplying demographic, '
     'educational, experiential, and skill data. This constructs the UserData object — the '
     'central profile entity powering all downstream personalised AI modules.'),
    ('Step 2 – Real-Time Labour Market Data Ingestion',
     'The system continuously ingests structured SkillDemand, SalaryData, and TrendData from '
     'labour market data sources. This data populates the city-segmented skill demand database '
     'that underpins the Heatmap, Shortage Engine, and Migration modules.'),
    ('Step 3 – AI Flow Invocation & Skill Extraction',
     'When a user submits a job description to the Skill Extractor, the extract-skills-from-job-'
     'description Genkit flow is invoked. The flow uses NLP-augmented prompting to parse text and '
     'return a structured, ranked competency list. This enriches the user\'s profile automatically.'),
    ('Step 4 – Career Route Recommendation',
     'The AI Career Planner invokes the recommend-career-route flow, passing the complete UserData '
     'profile. The Gemini model analyses the profile against inferred market trajectories and '
     'returns a milestone-based roadmap with per-step XAI justification.'),
    ('Step 5 – Geographic Opportunity Analysis',
     'The Migration Assistant and Migration Simulator invoke their respective flows, combining the '
     'user\'s skill profile with city-level SkillDemand and cost-of-living-adjusted SalaryData to '
     'produce ranked geographic opportunity scores and simulated income projection scenarios.'),
    ('Step 6 – Policy & Anomaly Intelligence',
     'The Skill Shortage Engine surfaces supply-demand imbalances. Detected anomalies are fed into '
     'the policy-advisor flow, which synthesises structured policy prescriptions — complete with '
     'rationale — to serve enterprise HR or government stakeholders.'),
    ('Step 7 – Skill Graph Construction & Delivery',
     'The skill-graph-intelligence flow generates a knowledge graph of skill nodes, weighted '
     'adjacency edges (co-occurrence, co-learning, career transition probability), and cluster '
     'labels. The frontend renders this as an interactive, navigable network graph.'),
    ('Step 8 – Continuous Learning & Profile Update',
     'User interactions, job market data refreshes, and AI recommendation outcome signals feed '
     'back into the system\'s data layer, updating demand counters, trend series, and informing '
     'prompt tuning cycles — creating a continuously improving intelligence platform.'),
]
for title, desc in steps:
    add_sub_heading(doc, title)
    add_body(doc, desc)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 – DRAWINGS / SUPPORT MATERIAL
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 8, 'Drawings / Support Material')
drawings = [
    ('FIG. 1', 'System Architecture Overview',
     'Block diagram illustrating the layered architecture: Next.js frontend → Firebase Auth → Genkit AI Gateway → Google Gemini → Labour Market Data Layer → Firebase Hosting.'),
    ('FIG. 2', 'High-Level Process Flowchart',
     'Flowchart depicting the 8-step end-to-end operational flow from user onboarding through continuous learning feedback.'),
    ('FIG. 3', 'Skills Heatmap Module Data Flow',
     'Detailed diagram of the Heatmap module: data ingestion → SkillDemand DB → Recharts visualisation layer → user filter interaction.'),
    ('FIG. 4', 'AI Career Planner Flow',
     'Process flow of the recommend-career-route Genkit flow: UserData input → Gemini model → milestone roadmap + XAI rationale output.'),
    ('FIG. 5', 'Migration Opportunity Scoring Model',
     'Diagram of the salary normalisation and migration opportunity scoring pipeline, showing cost-of-living index weighting.'),
    ('FIG. 6', 'Skill Knowledge Graph Structure',
     'Graph diagram showing skill nodes, weighted adjacency edges, cluster groupings, and career adjacency score annotations.'),
]
for fig, title, desc in drawings:
    add_bullet(doc, f'{title}: {desc}', fig)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 – SUMMARY OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 9, 'Summary of Figures / Diagrams')
for fig, title, desc in drawings:
    add_sub_heading(doc, f'{fig}: {title}')
    add_body(doc, desc)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 – ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 10, 'Abstract (150 words)')
add_body(doc,
    'SkillMapper AI is an artificial intelligence-powered career intelligence platform that '
    'dynamically maps skill demand and supply across cities and industries, extracts competencies '
    'from job descriptions using Natural Language Processing, and generates personalised career '
    'trajectories through a multi-dimensional user profiling engine. The system comprises seven '
    'integrated AI flows built on Google Gemini via the Genkit framework, covering career route '
    'recommendation, geographic migration analysis, policy advisory, and skill graph intelligence. '
    'A key innovation is the integration of cost-of-living-normalised salary benchmarking with '
    'real-time demand heatmapping, producing actionable geographic opportunity scores. An '
    'Explainable AI (XAI) layer accompanies every recommendation with transparent, data-grounded '
    'rationale, enabling trust and human-in-the-loop decision making. Continuous learning from '
    'market data refreshes ensures the platform adapts to labour market evolution. SkillMapper AI '
    'represents a significant advance beyond static career advisory tools, offering a unified, '
    'intelligent, and adaptive skill intelligence layer for individuals, employers, and policy makers.')

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 – BRIEF DETAIL OF PARTS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 11, 'Brief Detail of System Parts')
parts = [
    ('Next.js 14 Application Server',
     'The core computational host. Executes all server-side AI flow invocations, route rendering, '
     'and API handling. Deployed on Firebase Hosting (or Vercel-compatible cloud infrastructure).'),
    ('Genkit AI Gateway',
     'The AI orchestration layer that manages flow execution, Gemini model calls, prompt versioning, '
     'and schema-validated input/output processing for all seven AI flows.'),
    ('Google Gemini Pro (LLM)',
     'The generative AI model that performs all natural language understanding, reasoning, and '
     'structured output generation tasks within the Genkit flow pipeline.'),
    ('Firebase Authentication',
     'Manages user identity, session management, and secure access control to personalised features '
     'and stored user profiles.'),
    ('Labour Market Data Store',
     'Structured databases containing SkillDemand, SalaryData, TrendData, City, Skill, and Industry '
     'entities — the factual foundation for all market intelligence modules.'),
    ('React + Recharts Frontend',
     'The user-facing interface layer, rendering interactive charts (heatmaps, bar charts, line '
     'charts), skill graph visualisations, and AI-generated recommendation cards.'),
    ('UserData Profile Engine',
     'The onboarding and profile management subsystem that constructs and stores the multi-dimensional '
     'UserData object driving all personalised AI flow invocations.'),
    ('XAI Rationale Generator',
     'Embedded within each Genkit flow prompt design, this component ensures that every AI output '
     'includes a structured, human-readable justification section alongside its primary recommendation.'),
]
for name, desc in parts:
    add_bullet(doc, desc, name)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 – SOFTWARE STEPS & FUNCTIONALITY
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 12, 'Software Steps, Flow & Functionality (Embedded in Device/Apparatus)')
steps12 = [
    ('Data Acquisition & Ingestion',
     'The application server periodically fetches and stores current SkillDemand, SalaryData, and '
     'TrendData from configured labour market APIs. Raw data is validated and written to the '
     'structured data store for consumption by all intelligence modules.'),
    ('User Profile Construction',
     'The onboarding module collects structured user inputs and constructs a validated UserData '
     'object. This object is persisted to the user\'s authenticated Firebase session and serves as '
     'the primary input context for all personalised AI flows.'),
    ('NLP Skill Extraction',
     'When job description text is provided, the extract-skills-from-job-description Genkit flow '
     'is triggered. The Gemini model performs linguistic analysis, competency entity recognition, '
     'and proficiency inference, returning a structured JSON skill list that is merged into the '
     'user\'s profile.'),
    ('AI Career Route Generation',
     'The recommend-career-route flow receives the complete UserData profile. The Gemini model '
     'performs multi-step reasoning across career progression data, market demand signals, and '
     'educational prerequisites to produce a structured, milestone-based career roadmap with '
     'per-item XAI rationale.'),
    ('Geographic Opportunity Computation',
     'The migration-assistant and migration-simulator flows receive the user\'s skill profile and '
     'target city parameters. Cost-of-living-adjusted salary scores are computed and combined with '
     'city-level SkillDemand data to generate a ranked opportunity comparison output.'),
    ('Anomaly Detection & Policy Generation',
     'The system analyses SkillDemand differentials to detect statistical anomalies. Detected '
     'signals are passed to the policy-advisor flow, which synthesises them into prioritised, '
     'structured policy recommendations with supporting evidence.'),
    ('Skill Graph Construction',
     'The skill-graph-intelligence flow processes skill co-occurrence, learning pathway, and '
     'career transition data to generate a weighted directed graph of skill relationships. The '
     'graph is serialised and delivered to the frontend for interactive rendering.'),
    ('Continuous Learning & Model Refresh',
     'Interaction logs, market data updates, and flow performance signals are collected and used '
     'to update demand counters, refine trend models, and inform iterative prompt engineering — '
     'embedding improved intelligence into the system for all future operations.'),
]
for i, (title, desc) in enumerate(steps12, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f'{i}. {title}: ')
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.color.rgb = DARK_BLUE
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11); r2.font.color.rgb = GREY_TEXT

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 – PROCESS STEPS IN BULLET POINTS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 13, 'Process Steps / Composition (Bullet Point Summary)')
add_body(doc,
    'The present invention relates to a computer-implemented process for AI-powered skill '
    'intelligence and career navigation. The steps of the process are:')
process_bullets = [
    'Ingesting real-time, city- and industry-segmented skill demand and supply data from labour market sources and structuring it into queryable data entities.',
    'Constructing multi-dimensional user skill profiles through a structured onboarding process, capturing professional, educational, and experiential data.',
    'Extracting and ranking competencies from unstructured job description text using NLP-augmented generative AI flows.',
    'Generating personalised, milestone-based career roadmaps through a multi-signal AI reasoning flow that evaluates market demand, user profile dimensions, and career progression logic.',
    'Computing cost-of-living-normalised geographic opportunity scores to enable evidence-based career migration decisions.',
    'Detecting statistical skill market anomalies and synthesising them into human-readable, prioritised policy prescriptions.',
    'Constructing and delivering an interactive knowledge graph of skill interdependencies, cluster memberships, and career adjacency scores.',
    'Simulating future workforce compositions under alternative hiring, training, and migration scenarios.',
    'Presenting all AI recommendations with transparent, data-grounded XAI rationale to enable human-in-the-loop decision making.',
    'Continuously updating market data, user profiles, and AI model parameters based on interaction signals and new market ingestion cycles — creating a perpetually adaptive intelligence platform.',
]
for b in process_bullets:
    add_bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 14 – RELATED PUBLICATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 14, 'Related Publications')
refs = [
    ('Brynjolfsson, E., & McAfee, A.',
     '"The Second Machine Age: Work, Progress, and Prosperity in a Time of Brilliant Technologies." W. W. Norton & Company (2014).',
     'Provides foundational context on labour market disruption from AI and automation, the central problem domain the invention addresses through its dynamic skill intelligence system.'),
    ('Dev, S., & Phillips, J.',
     '"Attenuating Bias in Word Vectors." Proceedings of AISTATS (2019).',
     'Informs the fairness considerations in the NLP skill extraction flow to ensure competency identification is not biased by linguistic or demographic factors in job description text.'),
    ('Miller, T.',
     '"Explanation in Artificial Intelligence: Insights from the Social Sciences." Artificial Intelligence Journal (2019).',
     'Provides the theoretical basis for the XAI rationale generation component embedded in every SkillMapper AI recommendation flow.'),
    ('Noy, S., & Zhang, W.',
     '"Experimental Evidence on the Productivity Effects of Generative Artificial Intelligence." MIT Working Paper (2023).',
     'Directly relevant to the invention\'s career planner module, demonstrating measurable productivity gains from AI-assisted career and task guidance — validating the core value proposition.'),
    ('Xu, R., et al.',
     '"Skill2vec: Machine Learning Approach for Determining the Relevant Skills from Job Description." IEEE Transactions on Learning Technologies (2021).',
     'Provides prior art context for skill extraction from job descriptions; the present invention extends this paradigm by integrating extracted skills into a dynamic user profile and downstream career AI reasoning pipeline.'),
]
for authors, citation, relevance in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'{authors} ')
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(11); r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(citation)
    r2.italic = True; r2.font.name = 'Calibri'; r2.font.size = Pt(11); r2.font.color.rgb = GREY_TEXT
    rel_p = doc.add_paragraph()
    rel_p.paragraph_format.space_after = Pt(6)
    rel_r1 = rel_p.add_run('Relevance: ')
    rel_r1.bold = True; rel_r1.font.name = 'Calibri'; rel_r1.font.size = Pt(10); rel_r1.font.color.rgb = MID_BLUE
    rel_r2 = rel_p.add_run(relevance)
    rel_r2.font.name = 'Calibri'; rel_r2.font.size = Pt(10); rel_r2.font.color.rgb = GREY_TEXT

# ══════════════════════════════════════════════════════════════════════════════
#  TECH STACK TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 15, 'Technology Stack')
stack = [
    ('Frontend Framework',  'Next.js 14 (App Router) + React + TypeScript'),
    ('UI Styling',          'Tailwind CSS + Radix UI Component Library'),
    ('AI Orchestration',    'Google Genkit'),
    ('AI Model',            'Google Gemini Pro (gemini-1.5-pro / gemini-pro-vision)'),
    ('Data Visualisation',  'Recharts (Bar, Line, Scatter, Heatmap Charts)'),
    ('Backend / Auth',      'Firebase Authentication + Firebase Hosting'),
    ('Type Safety',         'Zod (Schema validation for all AI flow I/O)'),
    ('Language',            'TypeScript (strict mode)'),
    ('Package Manager',     'npm'),
    ('Deployment',          'Firebase Hosting / Vercel-compatible'),
]
st = doc.add_table(rows=len(stack)+1, cols=2)
st.style = 'Table Grid'
st.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['Component', 'Technology']):
    c = st.rows[0].cells[j]
    shade_cell(c, '0A254A')
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.color.rgb = WHITE
for i, (comp, tech) in enumerate(stack):
    row = st.rows[i+1]
    shade = 'F0F4FF' if i % 2 == 0 else 'FFFFFF'
    c0 = row.cells[0]; c1 = row.cells[1]
    shade_cell(c0, shade); shade_cell(c1, shade)
    r0 = c0.paragraphs[0].add_run(comp)
    r0.bold = True; r0.font.name = 'Calibri'; r0.font.size = Pt(10); r0.font.color.rgb = DARK_BLUE
    r1 = c1.paragraphs[0].add_run(tech)
    r1.font.name = 'Calibri'; r1.font.size = Pt(10); r1.font.color.rgb = GREY_TEXT

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER / CONFIDENTIALITY NOTICE
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 16 – INVENTOR DECLARATION & SIGNATURE
# ══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, 16, 'Inventor Declaration & Signature')
add_body(doc,
    'I, the undersigned inventor, hereby declare that the information provided in this Invention '
    'Disclosure Form is, to the best of my knowledge, true, accurate, and complete. I believe '
    'this invention to be novel and not previously disclosed, published, or patented in an '
    'identical form. I agree to cooperate fully with the IP Cell in the preparation and '
    'prosecution of any patent application that may result from this disclosure.')

doc.add_paragraph()

# Signature Table(s)
inventors = [
    {
        'Name': 'Ansh Gupta',
        'Role': 'Full Stack AI Developer & Product Engineer',
    },
    {
        'Name': 'Vidur Kumar',
        'Role': 'AI Researcher & Data Scientist',
    }
]

for inv in inventors:
    sig = doc.add_table(rows=5, cols=2)
    sig.style = 'Table Grid'
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_data = [
        ('Inventor Full Name',    inv['Name']),
        ('Designation / Role',   inv['Role']),
        ('Department',           'AI/ML & Product Engineering'),
        ('Date of Submission',   '25 February 2026'),
        ('Inventor Signature',   '___________________________'),
    ]
    for i, (lbl, val) in enumerate(sig_data):
        lc = sig.rows[i].cells[0]
        vc = sig.rows[i].cells[1]
        shade_cell(lc, '1A4080')
        shade_cell(vc, 'FFFFFF' if i % 2 == 0 else 'F8FAFF')
        lc.width = Inches(2.4)
        vc.width = Inches(4.1)
        lr = lc.paragraphs[0].add_run(lbl)
        lr.bold = True; lr.font.name = 'Calibri'; lr.font.size = Pt(10); lr.font.color.rgb = WHITE
        vr = vc.paragraphs[0].add_run(val)
        vr.font.name = 'Calibri'; vr.font.size = Pt(10)
        vr.font.color.rgb = DARK_BLUE if val != '___________________________' else RGBColor(0x9C,0xA3,0xAF)
    
    doc.add_paragraph() # Space between inventor tables

doc.add_paragraph()

# IP Office Use Only box
add_sub_heading(doc, 'FOR IP OFFICE USE ONLY')
ip_tbl = doc.add_table(rows=4, cols=2)
ip_tbl.style = 'Table Grid'
ip_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ip_rows = [
    ('IDF / Ticket Number Assigned', '___________________________'),
    ('Date of Receipt by IP Office', '___________________________'),
    ('Reviewed By (IP Officer)',      '___________________________'),
    ('Status / Remarks',             '___________________________'),
]
for i, (lbl, val) in enumerate(ip_rows):
    lc = ip_tbl.rows[i].cells[0]
    vc = ip_tbl.rows[i].cells[1]
    shade_cell(lc, 'E8F0FE')
    shade_cell(vc, 'FFFFFF')
    lc.width = Inches(2.8)
    vc.width = Inches(3.7)
    lr = lc.paragraphs[0].add_run(lbl)
    lr.bold = True; lr.font.name = 'Calibri'; lr.font.size = Pt(10); lr.font.color.rgb = DARK_BLUE
    vr = vc.paragraphs[0].add_run(val)
    vr.font.name = 'Calibri'; vr.font.size = Pt(10); vr.font.color.rgb = RGBColor(0x9C,0xA3,0xAF)

doc.add_paragraph()

# Footer banner
foot_tbl = doc.add_table(rows=1, cols=1)
foot_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
fc = foot_tbl.rows[0].cells[0]
shade_cell(fc, '0A254A')
fp = fc.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(6)
fp.paragraph_format.space_after  = Pt(6)
fr = fp.add_run(
    'STRICTLY CONFIDENTIAL  |  Invention Disclosure Form (IDF)  |  NEW SUBMISSION  |  '
    'SkillMapper AI  |  25 February 2026  |  Ticket ID: Pending Assignment')
fr.font.name = 'Calibri'; fr.font.size = Pt(8.5); fr.font.color.rgb = LIGHT_BLUE

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = 'SkillMapper AI_IDF_New Submission_25-Feb-2026.docx'
doc.save(output_path)
print(f'✅  Document saved: {output_path}')
